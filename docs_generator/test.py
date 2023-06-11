from docx import Document
from docx.shared import Inches

document = Document('test.docx')

# Находим параграф, содержащий "PLACEHOLDER"
for p in document.paragraphs:
    if "ЛИЧНАЯ ПОДПИСЬ" in p.text:
        # Удаляем текст "PLACEHOLDER"
        p.text = ""

        # Вставляем изображение вместо текста
        image_path = "posrt.jpeg"
        paragraph = p.insert_paragraph_before("")
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(1.25))
        break

document.save('test.docx')
